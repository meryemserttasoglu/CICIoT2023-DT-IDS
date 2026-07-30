"""
=============================================================================
MATEMATİKSEL MODEL DETAYLARI — Word Belgesi Oluşturucu
=============================================================================
Hybrid Digital Twin IDS sistemindeki TÜM matematiksel formülasyonları
detaylı olarak dokümante eder.
=============================================================================
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── STYLES ──────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    hs.font.bold = True

def add_formula(doc, formula_text, label=""):
    """Formülü kutu içinde göster."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(formula_text)
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
    run.bold = True
    if label:
        run2 = p.add_run(f"    ({label})")
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run2.font.name = "Calibri"

def add_note(doc, text):
    """Açıklama notu ekle."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run("📌 " + text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

def add_table_row(table, cells_data, bold=False, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        if bold or header:
            run.bold = True
        if header:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def make_simple_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Calibri"
    # Data
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Calibri"
    doc.add_paragraph()  # spacing
    return table


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  BAŞLIK                                                                  ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

title = doc.add_heading("", level=0)
run = title.add_run("CICIoT2023 Hybrid Digital Twin IDS\nMatematiksel Model Detayları")
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Tüm Pipeline Bileşenlerinin Formülasyonu ve Matematiksel Temelleri")
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()  # boşluk

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  1. ÖN İŞLEME — StandardScaler                                          ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

doc.add_heading("1. Veri Ön İşleme — Normalizasyon (StandardScaler)", level=1)

doc.add_paragraph(
    "Tüm sayısal özellikler StandardScaler ile normalize edilir. Bu dönüşüm, "
    "her özelliğin ortalamasını 0, standart sapmasını 1 yapar. Böylece farklı "
    "ölçeklerdeki özellikler (ör. flow_duration vs. flag_count) eşit ağırlıkta "
    "değerlendirilir."
)

doc.add_heading("1.1 StandardScaler Formülü", level=2)

add_formula(doc, "x'ᵢ = (xᵢ − μᵢ) / σᵢ", "Eq.1")

doc.add_paragraph("Burada:")
items = [
    ("xᵢ", "i. özelliğin ham değeri"),
    ("μᵢ", "i. özelliğin eğitim setindeki ortalaması: μᵢ = (1/N) Σⱼ xⱼᵢ"),
    ("σᵢ", "i. özelliğin eğitim setindeki standart sapması: σᵢ = √[(1/N) Σⱼ (xⱼᵢ − μᵢ)²]"),
    ("x'ᵢ", "normalize edilmiş değer (z-score)"),
]
for sym, desc in items:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(sym)
    run.bold = True
    run.font.name = "Consolas"
    p.add_run(f" : {desc}")

add_note(doc, "Veri sızıntısını önlemek için scaler YALNIZCA eğitim verisi üzerinde fit edilir, "
         "sonra hem eğitim hem test verisine transform uygulanır.")

doc.add_heading("1.2 Etiket Kodlama (Label Encoding)", level=2)

doc.add_paragraph(
    "Üç seviyede etiket kodlama uygulanır:"
)

make_simple_table(doc,
    ["Seviye", "Sınıf Sayısı", "Kodlama", "Kullanım"],
    [
        ["İkili (Binary)", "2", "y = 0 (Benign), 1 (Attack)", "AE anomali tespiti"],
        ["Kategori", "8", "LabelEncoder → {0..7}", "ConvLSTM çok-sınıf"],
        ["Saldırı Tipi", "34", "LabelEncoder → {0..33}", "İnce taneli analiz"],
    ]
)

doc.add_heading("1.3 Veri Bölme Stratejisi", level=2)

add_formula(doc, "D = D_train (80%) ∪ D_test (20%),  stratified by category", "Eq.2")

add_note(doc, "Sınıf dengesizliğini korumak için stratified sampling uygulanır. "
         "Orijinal veri setinden %10 örneklem alınır (SAMPLE_FRACTION = 0.10).")


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  2. AUTOENCODER — DİJİTAL İKİZ                                          ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

doc.add_page_break()
doc.add_heading("2. Autoencoder — Dijital İkiz Modeli", level=1)

doc.add_paragraph(
    "Autoencoder (AE), sistemin Dijital İkiz bileşenidir. Ağ SADECE normal (benign) "
    "trafik üzerinde eğitilerek, normal ağ davranışının matematiksel bir temsilini "
    "(dijital ikizini) öğrenir. Saldırı trafiği AE tarafından doğru şekilde "
    "rekonstrükte edilemez ve yüksek hata üretir."
)

doc.add_heading("2.1 Mimari Tanım", level=2)

doc.add_paragraph(
    "AE, simetrik bir encoder-decoder yapısına sahiptir. Encoder girişi sıkıştırır, "
    "decoder sıkıştırılmış temsilden orijinali yeniden oluşturur."
)

add_formula(doc, "Encoder:  f_enc : ℝ⁴⁶ → ℝ¹⁶", "Eq.3")
add_formula(doc, "Decoder:  f_dec : ℝ¹⁶ → ℝ⁴⁶", "Eq.4")
add_formula(doc, "Autoencoder:  f_AE(x) = f_dec(f_enc(x)) ≈ x", "Eq.5")

doc.add_heading("2.2 Katman Yapısı", level=2)

make_simple_table(doc,
    ["Katman", "Boyut", "Aktivasyon", "Ek Bileşenler"],
    [
        ["Giriş (Input)", "46", "—", "—"],
        ["Encoder Dense 1", "64", "ReLU", "BatchNorm + Dropout(0.2)"],
        ["Encoder Dense 2", "32", "ReLU", "BatchNorm + Dropout(0.2)"],
        ["Bottleneck (Darboğaz)", "16", "ReLU", "Latent temsil z"],
        ["Decoder Dense 1", "32", "ReLU", "BatchNorm + Dropout(0.2)"],
        ["Decoder Dense 2", "64", "ReLU", "BatchNorm + Dropout(0.2)"],
        ["Çıkış (Output)", "46", "Linear", "Rekonstrüksiyon x̂"],
    ]
)

doc.add_heading("2.3 Encoder Formülasyonu", level=2)

doc.add_paragraph(
    "Encoder, girişi ardışık lineer dönüşüm + non-lineer aktivasyon + "
    "normalizasyon katmanlarından geçirir:"
)

add_formula(doc, "h₁ = Dropout(BN(ReLU(W₁x + b₁))),  W₁ ∈ ℝ⁶⁴ˣ⁴⁶", "Eq.6")
add_formula(doc, "h₂ = Dropout(BN(ReLU(W₂h₁ + b₂))),  W₂ ∈ ℝ³²ˣ⁶⁴", "Eq.7")
add_formula(doc, "z = ReLU(W₃h₂ + b₃),  W₃ ∈ ℝ¹⁶ˣ³²  ← Bottleneck vektörü", "Eq.8")

doc.add_paragraph("Burada:")
items2 = [
    ("ReLU(a)", "max(0, a) — Doğrusal olmayan aktivasyon fonksiyonu"),
    ("BN(a)", "Batch Normalization: (a − μ_batch) / σ_batch × γ + β"),
    ("Dropout(a, p=0.2)", "Eğitimde rastgele %20 nöronu sıfırlama (regülarizasyon)"),
    ("z ∈ ℝ¹⁶", "Latent temsil — girişin sıkıştırılmış dijital ikizi"),
]
for sym, desc in items2:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(sym)
    run.bold = True
    run.font.name = "Consolas"
    p.add_run(f" : {desc}")

doc.add_heading("2.4 Decoder Formülasyonu", level=2)

doc.add_paragraph("Decoder, encoder'ın simetrik aynasıdır:")

add_formula(doc, "h₃ = Dropout(BN(ReLU(W₄z + b₄))),  W₄ ∈ ℝ³²ˣ¹⁶", "Eq.9")
add_formula(doc, "h₄ = Dropout(BN(ReLU(W₅h₃ + b₅))),  W₅ ∈ ℝ⁶⁴ˣ³²", "Eq.10")
add_formula(doc, "x̂ = W₆h₄ + b₆,  W₆ ∈ ℝ⁴⁶ˣ⁶⁴  ← Rekonstrüksiyon çıkışı", "Eq.11")

add_note(doc, "Çıkış katmanında linear aktivasyon kullanılır çünkü normalize "
         "edilmiş girdi değerleri (-∞, +∞) aralığındadır.")

doc.add_heading("2.5 Kayıp Fonksiyonu — Mean Squared Error (MSE)", level=2)

doc.add_paragraph(
    "AE, giriş ile çıkış arasındaki ortalama karesel hatayı minimize eder. "
    "Bu kayıp fonksiyonu, modelin normal trafiğin dijital ikizini ne kadar "
    "iyi öğrendiğini ölçer."
)

add_formula(doc, "L_MSE = (1/N) Σᵢ₌₁ᴺ (1/d) Σⱼ₌₁ᵈ (xᵢⱼ − x̂ᵢⱼ)²", "Eq.12")

doc.add_paragraph("Burada:")
items3 = [
    ("N", "Mini-batch'teki örnek sayısı (batch_size = 256)"),
    ("d", "Özellik boyutu (d = 46)"),
    ("xᵢⱼ", "i. örneğin j. özelliğinin orijinal değeri"),
    ("x̂ᵢⱼ", "i. örneğin j. özelliğinin rekonstrükte edilmiş değeri"),
]
for sym, desc in items3:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(sym)
    run.bold = True
    run.font.name = "Consolas"
    p.add_run(f" : {desc}")

doc.add_heading("2.6 Optimizasyon — Adam Optimizer", level=2)

doc.add_paragraph(
    "Adam (Adaptive Moment Estimation) optimizer kullanılır. Hem gradyanın "
    "birinci momenti (ortalama) hem ikinci momenti (varyans) takip edilir:"
)

add_formula(doc, "mₜ = β₁ · mₜ₋₁ + (1 − β₁) · gₜ           (1. moment tahmini)", "Eq.13")
add_formula(doc, "vₜ = β₂ · vₜ₋₁ + (1 − β₂) · gₜ²          (2. moment tahmini)", "Eq.14")
add_formula(doc, "m̂ₜ = mₜ / (1 − β₁ᵗ),  v̂ₜ = vₜ / (1 − β₂ᵗ)  (bias düzeltme)", "Eq.15")
add_formula(doc, "θₜ = θₜ₋₁ − α · m̂ₜ / (√v̂ₜ + ε)", "Eq.16")

make_simple_table(doc,
    ["Parametre", "Değer", "Açıklama"],
    [
        ["α (learning_rate)", "1×10⁻³", "Başlangıç öğrenme oranı"],
        ["β₁", "0.9", "1. moment bozunma katsayısı"],
        ["β₂", "0.999", "2. moment bozunma katsayısı"],
        ["ε", "1×10⁻⁷", "Sayısal kararlılık sabiti"],
    ]
)

doc.add_heading("2.7 Eğitim Düzenleme (Callbacks)", level=2)

make_simple_table(doc,
    ["Callback", "Parametre", "Açıklama"],
    [
        ["EarlyStopping", "patience = 7", "val_loss 7 epoch iyileşmezse dur"],
        ["ReduceLROnPlateau", "factor = 0.5, patience = 3", "val_loss 3 epoch iyileşmezse lr'ı yarıya düşür"],
    ]
)

add_note(doc, "Eğitim SADECE benign (normal) trafik ile yapılır. Saldırı örnekleri "
         "eğitim sırasında asla kullanılmaz — bu, Dijital İkiz konseptinin temelidir.")


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  3. ANOMALİ TESPİTİ — Eşik Hesaplama                                   ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

doc.add_page_break()
doc.add_heading("3. Anomali Tespiti — Eşik Tabanlı Sınıflandırma", level=1)

doc.add_paragraph(
    "Eğitilmiş AE, test verisindeki her örnek için bir rekonstrüksiyon hatası "
    "üretir. Bu hata, örneğin normal davranıştan ne kadar saptığını gösterir."
)

doc.add_heading("3.1 Rekonstrüksiyon Hatası (Anomali Skoru)", level=2)

doc.add_paragraph("Her test örneği xᵢ için anomali skoru:")

add_formula(doc, "RE(xᵢ) = (1/d) Σⱼ₌₁ᵈ (xᵢⱼ − x̂ᵢⱼ)²", "Eq.17")

doc.add_paragraph(
    "Bu, her özellik boyutundaki hataların ortalamasıdır — yüksek RE → "
    "normal davranıştan sapma → potansiyel saldırı."
)

doc.add_heading("3.2 Özellik Bazında Rekonstrüksiyon Hatası", level=2)

add_formula(doc, "RE_feature(xᵢ, j) = (xᵢⱼ − x̂ᵢⱼ)²,  ∀j ∈ {1, ..., 46}", "Eq.18")

add_note(doc, "Bu değer Aşama 7'deki Kök Neden Analizi'nde kullanılır. "
         "Hangi özelliğin en çok saptığını belirler.")

doc.add_heading("3.3 Eşik Değeri Belirleme", level=2)

doc.add_paragraph(
    "Anomali eşiği, benign (normal) test örneklerinin rekonstrüksiyon "
    "hatalarının 95. yüzdeliği olarak belirlenir:"
)

add_formula(doc, "τ = P₉₅(RE_benign) = Q₀.₉₅({RE(xᵢ) | yᵢ = 0})", "Eq.19")

doc.add_paragraph("Sınıflandırma kuralı:")

add_formula(doc, "ŷᵢ = { 1 (Saldırı)  eğer RE(xᵢ) > τ", "Eq.20a")
add_formula(doc, "       { 0 (Normal)   eğer RE(xᵢ) ≤ τ", "Eq.20b")

doc.add_heading("3.4 Eşik Duyarlılık Analizi", level=2)

doc.add_paragraph(
    "Farklı yüzdelik dilimlerle eşik hassasiyeti incelenir:"
)

make_simple_table(doc,
    ["Yüzdelik", "Etki", "Yorum"],
    [
        ["P₉₀", "Daha hassas (düşük eşik)", "Daha çok saldırı yakalar, daha çok yanlış alarm"],
        ["P₉₂", "Orta-hassas", "Dengeli yaklaşım"],
        ["P₉₅ ◄", "Varsayılan", "Optimal F1 dengesi"],
        ["P₉₇", "Daha tutucu (yüksek eşik)", "Az yanlış alarm, bazı saldırıları kaçırabilir"],
        ["P₉₉", "Çok tutucu", "Sadece en belirgin anomalileri yakalar"],
    ]
)

doc.add_heading("3.5 ROC-AUC Hesaplama", level=2)

add_formula(doc, "AUC = ∫₀¹ TPR(t) d(FPR(t))", "Eq.21")

doc.add_paragraph("Burada:")
for sym, desc in [("TPR", "True Positive Rate = TP / (TP + FN)"),
                   ("FPR", "False Positive Rate = FP / (FP + TN)"),
                   ("t", "Eşik değeri (τ) sürekli değiştirilerek hesaplanır")]:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(sym)
    run.bold = True
    run.font.name = "Consolas"
    p.add_run(f" : {desc}")

add_note(doc, "ROC-AUC < 0.5 ise anomali yönü ters çevrilir: anomaly_scores = −RE "
         "kullanılır. Bu, AE'nin henüz benign manifoldu iyi öğrenemediği durumlarda gerçekleşir.")


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  4. Conv1D-LSTM — Zamansal Model                                        ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

doc.add_page_break()
doc.add_heading("4. Conv1D-LSTM — Zamansal Örüntü Modeli", level=1)

doc.add_paragraph(
    "Conv1D-LSTM, ağ trafiğinin zamansal dinamiklerini yakalayan ikinci ana modeldir. "
    "Kayar pencere (sliding window) ile oluşturulan diziler üzerinde "
    "yerel örüntü çıkarma (Conv1D) ve zamansal bağımlılık modelleme (LSTM) birleştirilir."
)

doc.add_heading("4.1 Kayar Pencere Dönüşümü", level=2)

add_formula(doc, "X_seq[i] = [x_{i}, x_{i+1}, ..., x_{i+w-1}] ∈ ℝ^{w × d}", "Eq.22")
add_formula(doc, "y_seq[i] = y_{i+w-1}  (son zaman adımının etiketi)", "Eq.23")

make_simple_table(doc,
    ["Parametre", "Değer", "Açıklama"],
    [
        ["w (window_size)", "10", "Pencere genişliği (zaman adımı sayısı)"],
        ["d (n_features)", "46", "Her zaman adımındaki özellik sayısı"],
        ["N_seq", "N − w + 1", "Toplam oluşan dizi sayısı"],
    ]
)

doc.add_heading("4.2 Conv1D Katmanları — Yerel Örüntü Çıkarma", level=2)

doc.add_paragraph(
    "1D Konvolüsyon, her pencere içindeki ardışık zaman adımları arasında "
    "yerel örüntüleri (spike, pattern) tespit eder:"
)

add_formula(doc, "Conv1D: y[t] = ReLU(Σₖ₌₀^{K-1} W_k · x[t+k] + b)", "Eq.24")

make_simple_table(doc,
    ["Katman", "Filtre Sayısı", "Kernel Boyutu", "Padding", "Çıkış Boyutu"],
    [
        ["Conv1D_1", "64", "3", "same", "(w, 64)"],
        ["BatchNorm_1", "—", "—", "—", "(w, 64)"],
        ["Conv1D_2", "32", "3", "same", "(w, 32)"],
        ["BatchNorm_2", "—", "—", "—", "(w, 32)"],
    ]
)

doc.add_heading("4.3 LSTM Katmanları — Zamansal Bağımlılık", level=2)

doc.add_paragraph(
    "LSTM (Long Short-Term Memory), uzun vadeli bağımlılıkları bir hücre durumu "
    "(cell state) aracılığıyla modeller. Her zaman adımında dört kapı işlemi gerçekleşir:"
)

add_formula(doc, "fₜ = σ(W_f · [hₜ₋₁, xₜ] + b_f)         (Forget Gate — Unutma Kapısı)", "Eq.25")
add_formula(doc, "iₜ = σ(W_i · [hₜ₋₁, xₜ] + b_i)         (Input Gate — Girdi Kapısı)", "Eq.26")
add_formula(doc, "C̃ₜ = tanh(W_C · [hₜ₋₁, xₜ] + b_C)     (Aday Hücre Durumu)", "Eq.27")
add_formula(doc, "Cₜ = fₜ ⊙ Cₜ₋₁ + iₜ ⊙ C̃ₜ              (Hücre Durumu Güncelleme)", "Eq.28")
add_formula(doc, "oₜ = σ(W_o · [hₜ₋₁, xₜ] + b_o)         (Output Gate — Çıktı Kapısı)", "Eq.29")
add_formula(doc, "hₜ = oₜ ⊙ tanh(Cₜ)                       (Gizli Durum)", "Eq.30")

doc.add_paragraph("Burada:")
for sym, desc in [
    ("σ", "Sigmoid fonksiyonu: σ(x) = 1 / (1 + e⁻ˣ), çıkış ∈ (0, 1)"),
    ("⊙", "Eleman bazında çarpım (Hadamard product)"),
    ("fₜ", "Forget gate — önceki bilginin ne kadarının unutulacağı"),
    ("iₜ", "Input gate — yeni bilginin ne kadarının ekleneceği"),
    ("oₜ", "Output gate — hücre durumunun ne kadarının çıktıya yansıyacağı"),
    ("Cₜ", "Cell state — modelin uzun vadeli hafızası"),
    ("hₜ", "Hidden state — modelin kısa vadeli çıktısı"),
]:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(sym)
    run.bold = True
    run.font.name = "Consolas"
    p.add_run(f" : {desc}")

make_simple_table(doc,
    ["LSTM Katmanı", "Birim Sayısı", "Return Sequences", "Dropout"],
    [
        ["LSTM_1", "64", "True (dizi çıkışı)", "0.3"],
        ["LSTM_2", "32", "False (tek vektör)", "0.3"],
    ]
)

doc.add_heading("4.4 Sınıflandırma Başlığı", level=2)

add_formula(doc, "h_fc = ReLU(W_fc · h_lstm2 + b_fc),  W_fc ∈ ℝ⁶⁴ˣ³²", "Eq.31")
add_formula(doc, "ŷ = Softmax(W_out · h_fc + b_out),  W_out ∈ ℝ⁸ˣ⁶⁴", "Eq.32")

doc.add_heading("4.5 Softmax Fonksiyonu", level=2)

add_formula(doc, "Softmax(zⱼ) = e^{zⱼ} / Σₖ₌₁ᴷ e^{zₖ},  j = 1, ..., K", "Eq.33")

add_note(doc, "K = 8 sınıf (Benign, DDoS, DoS, Mirai, Recon, Spoofing, Web, BruteForce). "
         "Softmax çıkışı, her sınıfa ait olasılık vektörüdür: p ∈ ℝ⁸, Σpₖ = 1")

doc.add_heading("4.6 Kayıp Fonksiyonu — Categorical Cross-Entropy", level=2)

add_formula(doc, "L_CCE = −(1/N) Σᵢ₌₁ᴺ Σₖ₌₁ᴷ yᵢₖ · log(ŷᵢₖ)", "Eq.34")

doc.add_paragraph("Burada:")
for sym, desc in [
    ("yᵢₖ", "i. örneğin k. sınıfa ait one-hot etiketi (0 veya 1)"),
    ("ŷᵢₖ", "Softmax tarafından tahmin edilen olasılık"),
    ("K", "Toplam sınıf sayısı (8)"),
]:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(sym)
    run.bold = True
    run.font.name = "Consolas"
    p.add_run(f" : {desc}")

doc.add_heading("4.7 Sınıf Ağırlıklı Eğitim", level=2)

doc.add_paragraph("CICIoT2023'teki sınıf dengesizliğini ele almak için balanced class weights kullanılır:")

add_formula(doc, "wₖ = N / (K × nₖ)", "Eq.35")

doc.add_paragraph("Burada nₖ, k. sınıfın eğitim setindeki örnek sayısıdır. "
                  "Az temsil edilen sınıflar daha yüksek ağırlık alır.")


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  5. HİBRİT FÜZYON MODELİ                                                ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

doc.add_page_break()
doc.add_heading("5. Hibrit Füzyon Modeli — Geç Birleştirme (Late Fusion)", level=1)

doc.add_paragraph(
    "Sistemin temel katkısı olan hibrit model, AE'nin anomali skorlarını ve "
    "Conv1D-LSTM'in sınıf olasılıklarını tek bir karar katmanında birleştirir. "
    "Bu strateji 'geç birleştirme' (late fusion) olarak adlandırılır çünkü "
    "her model bağımsız olarak çıktı üretir ve birleştirme son aşamada yapılır."
)

doc.add_heading("5.1 Füzyon Vektörü Oluşturma", level=2)

add_formula(doc, "v_fusion = [ RE_feat(x) ‖ RE_global(x) ‖ p_ConvLSTM(x) ]", "Eq.36")

doc.add_paragraph("Bileşenler:")

make_simple_table(doc,
    ["Bileşen", "Boyut", "Kaynak", "Açıklama"],
    [
        ["RE_feat(x)", "46", "Autoencoder", "Her özelliğin rekonstrüksiyon hatası: (xⱼ − x̂ⱼ)²"],
        ["RE_global(x)", "1", "Autoencoder", "Global anomali skoru: mean(RE_feat)"],
        ["p_ConvLSTM(x)", "8", "Conv1D-LSTM", "8 sınıf için Softmax olasılık vektörü"],
        ["v_fusion", "55", "Birleşim", "Toplam füzyon vektörü (46 + 1 + 8 = 55)"],
    ]
)

doc.add_heading("5.2 Füzyon MLP Mimarisi", level=2)

doc.add_paragraph("Birleştirilmiş vektör, 3 katmanlı bir MLP'den geçirilir:")

make_simple_table(doc,
    ["Katman", "Giriş → Çıkış", "Aktivasyon", "Düzenleme"],
    [
        ["Fusion Dense 1", "55 → 128", "ReLU", "BatchNorm + Dropout(0.3)"],
        ["Fusion Dense 2", "128 → 64", "ReLU", "BatchNorm + Dropout(0.3)"],
        ["Fusion Dense 3", "64 → 32", "ReLU", "Dropout(0.2)"],
        ["Output", "32 → 8", "Softmax", "—"],
    ]
)

doc.add_heading("5.3 Füzyon Matematiksel Formülasyonu", level=2)

add_formula(doc, "h₁^f = Dropout(BN(ReLU(W₁^f · v_fusion + b₁^f)))", "Eq.37")
add_formula(doc, "h₂^f = Dropout(BN(ReLU(W₂^f · h₁^f + b₂^f)))", "Eq.38")
add_formula(doc, "h₃^f = Dropout(ReLU(W₃^f · h₂^f + b₃^f))", "Eq.39")
add_formula(doc, "ŷ_hybrid = Softmax(W₄^f · h₃^f + b₄^f)", "Eq.40")

add_note(doc, "Alt modeller (AE ve ConvLSTM) dondurulmuştur — yalnızca füzyon katmanlarının "
         "ağırlıkları eğitilir. Bu, bilgi kaybını önler ve eğitim süresini kısaltır.")

doc.add_heading("5.4 İkili Karar Dönüşümü", level=2)

doc.add_paragraph("Çok-sınıflı çıktıdan ikili saldırı tespiti:")

add_formula(doc, "ŷ_binary = { 0  eğer argmax(ŷ_hybrid) = idx_Benign", "Eq.41a")
add_formula(doc, "            { 1  aksi halde (saldırı)", "Eq.41b")


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  6. KÖK NEDEN ANALİZİ (RCA)                                             ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

doc.add_page_break()
doc.add_heading("6. Kök Neden Analizi (Root Cause Analysis — RCA)", level=1)

doc.add_paragraph(
    "RCA, tespit edilen anomalilerin NEDEN anomali olarak işaretlendiğini açıklar. "
    "İki bağımsız yöntem kullanılır ve sonuçları birleştirilir."
)

doc.add_heading("6.1 Yöntem 1: AE Tabanlı RCA — Özellik Bazında Hata", level=2)

doc.add_paragraph(
    "Her tespit edilen anomali için, hangi özelliğin en yüksek rekonstrüksiyon "
    "hatasına sahip olduğu belirlenir:"
)

add_formula(doc, "RCA_AE(j) = (1/N_anom) Σᵢ∈Anomalies (xᵢⱼ − x̂ᵢⱼ)²", "Eq.42")

add_formula(doc, "Top-K Root Causes = argsort(RCA_AE)[-K:]  (K = 5)", "Eq.43")

doc.add_paragraph(
    "İnterpretasyon: Yüksek RCA_AE(j) → j. özellik, normal davranıştan en çok "
    "sapan özellik → saldırının muhtemel nedenini gösterir."
)

doc.add_heading("6.2 Kategori Bazında RCA", level=2)

doc.add_paragraph(
    "Aynı formül, her saldırı kategorisi (DDoS, DoS, Mirai vb.) için ayrı ayrı hesaplanır:"
)

add_formula(doc, "RCA_AE^(c)(j) = (1/N_c) Σᵢ∈Category_c (xᵢⱼ − x̂ᵢⱼ)²", "Eq.44")

add_note(doc, "Bu, her saldırı tipinin kendine özgü 'parmak izini' ortaya çıkarır. "
         "Örneğin DDoS saldırıları syn_flag_number özelliğinde yüksek sapma gösterirken, "
         "Recon saldırıları Recon-PortScan için ack_count'ta sapma gösterebilir.")

doc.add_heading("6.3 Yöntem 2: SHAP Tabanlı RCA — Model Agnostik Açıklanabilirlik", level=2)

doc.add_paragraph(
    "SHAP (SHapley Additive exPlanations), oyun teorisinden türetilmiş "
    "Shapley değerlerini kullanarak her özelliğin model kararına katkısını ölçer."
)

add_formula(doc, "φⱼ = Σ_{S⊆F\\{j}} [ |S|! (|F|−|S|−1)! / |F|! ] × [f(S∪{j}) − f(S)]", "Eq.45")

doc.add_paragraph("Burada:")
for sym, desc in [
    ("φⱼ", "j. özelliğin Shapley değeri (model kararına katkısı)"),
    ("F", "Tüm özellikler kümesi (|F| = 46)"),
    ("S", "j hariç özellik alt kümesi"),
    ("f(S)", "Yalnızca S özelliklerini kullanarak modelin tahmini"),
    ("f(S∪{j})", "S'ye j özelliği eklendiğinde modelin tahmini"),
]:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(sym)
    run.bold = True
    run.font.name = "Consolas"
    p.add_run(f" : {desc}")

doc.add_paragraph(
    "TreeExplainer, Random Forest için bu hesaplamayı verimli şekilde yapar. "
    "Global özellik önemi, tüm örnekler üzerindeki ortalama mutlak SHAP değeridir:"
)

add_formula(doc, "Importance_SHAP(j) = (1/N) Σᵢ₌₁ᴺ |φⱼ(xᵢ)|", "Eq.46")

doc.add_heading("6.4 Birleşik RCA Skoru", level=2)

doc.add_paragraph(
    "İki yöntemin sonuçları normalize edilerek eşit ağırlıkla birleştirilir:"
)

add_formula(doc, "RCA_norm_AE(j) = RCA_AE(j) / max(RCA_AE)", "Eq.47")
add_formula(doc, "RCA_norm_SHAP(j) = Imp_SHAP(j) / max(Imp_SHAP)", "Eq.48")
add_formula(doc, "RCA_combined(j) = 0.5 × RCA_norm_AE(j) + 0.5 × RCA_norm_SHAP(j)", "Eq.49")


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  7. DEĞERLENDİRME METRİKLERİ                                            ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

doc.add_page_break()
doc.add_heading("7. Değerlendirme Metrikleri — Matematiksel Tanımlar", level=1)

doc.add_paragraph(
    "Sistem dört temel metrik ile değerlendirilir. Tüm metrikler confusion matrix "
    "(karışıklık matrisi) elemanlarından hesaplanır."
)

doc.add_heading("7.1 Karışıklık Matrisi", level=2)

make_simple_table(doc,
    ["", "Tahmin: Positive (Saldırı)", "Tahmin: Negative (Normal)"],
    [
        ["Gerçek: Positive", "TP (True Positive)", "FN (False Negative)"],
        ["Gerçek: Negative", "FP (False Positive)", "TN (True Negative)"],
    ]
)

doc.add_heading("7.2 Temel Metrikler", level=2)

add_formula(doc, "Accuracy = (TP + TN) / (TP + TN + FP + FN)", "Eq.50")
add_formula(doc, "Precision = TP / (TP + FP)", "Eq.51")
add_formula(doc, "Recall (Sensitivity) = TP / (TP + FN)", "Eq.52")
add_formula(doc, "F1-Score = 2 × (Precision × Recall) / (Precision + Recall)", "Eq.53")

doc.add_heading("7.3 Çok-sınıflı Metrikler (Weighted Average)", level=2)

doc.add_paragraph("8 sınıflı sınıflandırma için weighted average kullanılır:")

add_formula(doc, "Metric_weighted = Σₖ₌₁ᴷ (nₖ/N) × Metric_k", "Eq.54")

doc.add_paragraph("Burada nₖ, k. sınıfın test setindeki örnek sayısı, N toplam örnektir.")

doc.add_heading("7.4 ROC-AUC (Çok-sınıflı)", level=2)

add_formula(doc, "ROC-AUC_OVR = Σₖ₌₁ᴷ (nₖ/N) × AUC_k(One-vs-Rest)", "Eq.55")

doc.add_heading("7.5 Average Precision (AP)", level=2)

add_formula(doc, "AP = Σₙ (Rₙ − Rₙ₋₁) × Pₙ", "Eq.56")

doc.add_paragraph("Precision-Recall eğrisinin altındaki alan olarak hesaplanır.")


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  8. HİPERPARAMETRE ÖZETİ                                                ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

doc.add_page_break()
doc.add_heading("8. Hiperparametre Özet Tablosu", level=1)

doc.add_paragraph("Tüm modellerin hiperparametreleri config.py'den alınan değerlerdir:")

doc.add_heading("8.1 Autoencoder (Dijital İkiz)", level=2)

make_simple_table(doc,
    ["Parametre", "Sembol", "Değer"],
    [
        ["Giriş boyutu", "d", "46"],
        ["Encoder katmanları", "—", "[64, 32]"],
        ["Darboğaz boyutu", "z_dim", "16"],
        ["Decoder katmanları", "—", "[32, 64] (simetrik)"],
        ["Öğrenme oranı", "α", "1×10⁻³"],
        ["Batch boyutu", "B", "256"],
        ["Maksimum epoch", "E_max", "50"],
        ["Early stopping patience", "p_es", "7"],
        ["Doğrulama oranı", "—", "0.15 (%15)"],
        ["Dropout oranı", "p_drop", "0.2"],
        ["Kayıp fonksiyonu", "L", "MSE"],
    ]
)

doc.add_heading("8.2 Conv1D-LSTM (Zamansal Model)", level=2)

make_simple_table(doc,
    ["Parametre", "Sembol", "Değer"],
    [
        ["Pencere boyutu", "w", "10"],
        ["Conv1D filtre sayısı", "F", "64 → 32"],
        ["Conv1D kernel boyutu", "K", "3"],
        ["LSTM birimleri", "—", "64 → 32"],
        ["Öğrenme oranı", "α", "1×10⁻³"],
        ["Batch boyutu", "B", "128"],
        ["Maksimum epoch", "E_max", "30"],
        ["Early stopping patience", "p_es", "5"],
        ["Dropout oranı", "p_drop", "0.3 / 0.2"],
        ["Kayıp fonksiyonu", "L", "Categorical Cross-Entropy"],
        ["Sınıf ağırlıkları", "wₖ", "Balanced (otomatik)"],
    ]
)

doc.add_heading("8.3 Hibrit Füzyon Modeli", level=2)

make_simple_table(doc,
    ["Parametre", "Sembol", "Değer"],
    [
        ["Giriş boyutu", "d_fusion", "55 (46+1+8)"],
        ["Gizli katmanlar", "—", "128 → 64 → 32"],
        ["Öğrenme oranı", "α", "5×10⁻⁴"],
        ["Batch boyutu", "B", "128"],
        ["Maksimum epoch", "E_max", "30"],
        ["Early stopping patience", "p_es", "5"],
        ["Dropout oranları", "p_drop", "0.3, 0.3, 0.2"],
        ["Kayıp fonksiyonu", "L", "Categorical Cross-Entropy"],
    ]
)

doc.add_heading("8.4 Genel Pipeline Parametreleri", level=2)

make_simple_table(doc,
    ["Parametre", "Değer", "Açıklama"],
    [
        ["Veri örnekleme oranı", "0.10 (%10)", "Hız için rastgele alt örneklem"],
        ["Test oranı", "0.20 (%20)", "Stratified train/test bölme"],
        ["Random seed", "42", "Tekrarlanabilirlik"],
        ["Anomali eşiği yüzdeliği", "95", "P₉₅ benign hatası"],
        ["SHAP örneklem boyutu", "500", "Açıklanabilirlik analizi için"],
        ["RCA top-K", "5", "En önemli K kök neden özelliği"],
    ]
)


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  9. PARAMETRE SAYISI                                                     ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

doc.add_heading("9. Model Parametre Sayıları", level=1)

doc.add_paragraph(
    "Her modeldeki eğitilebilir parametre sayısı, katman boyutlarından hesaplanır:"
)

doc.add_heading("9.1 Autoencoder Parametre Hesabı", level=2)

make_simple_table(doc,
    ["Katman Geçişi", "Ağırlık (W)", "Bias (b)", "BN (γ, β)", "Toplam"],
    [
        ["46→64 (enc_1)", "46×64 = 2944", "64", "2×64 = 128", "3,136"],
        ["64→32 (enc_2)", "64×32 = 2048", "32", "2×32 = 64", "2,144"],
        ["32→16 (bottleneck)", "32×16 = 512", "16", "—", "528"],
        ["16→32 (dec_1)", "16×32 = 512", "32", "2×32 = 64", "608"],
        ["32→64 (dec_2)", "32×64 = 2048", "64", "2×64 = 128", "2,240"],
        ["64→46 (output)", "64×46 = 2944", "46", "—", "2,990"],
        ["TOPLAM", "", "", "", "≈ 11,646"],
    ]
)

doc.add_heading("9.2 Conv1D-LSTM Parametre Hesabı", level=2)

doc.add_paragraph("LSTM parametre formülü:")
add_formula(doc, "P_LSTM = 4 × [(n_input + n_units) × n_units + n_units]", "Eq.57")

make_simple_table(doc,
    ["Katman", "Hesaplama", "Parametre Sayısı"],
    [
        ["Conv1D_1 (46→64, k=3)", "3×46×64 + 64", "8,896"],
        ["Conv1D_2 (64→32, k=3)", "3×64×32 + 32", "6,176"],
        ["LSTM_1 (32→64)", "4×(32+64)×64 + 4×64", "24,832"],
        ["LSTM_2 (64→32)", "4×(64+32)×32 + 4×32", "12,416"],
        ["Dense FC (32→64)", "32×64 + 64", "2,112"],
        ["Output (64→8)", "64×8 + 8", "520"],
        ["+ BatchNorm", "—", "≈384"],
        ["TOPLAM", "", "≈ 55,336"],
    ]
)

doc.add_heading("9.3 Hibrit Füzyon Parametre Hesabı", level=2)

make_simple_table(doc,
    ["Katman", "Hesaplama", "Parametre Sayısı"],
    [
        ["Dense_1 (55→128)", "55×128 + 128", "7,168"],
        ["Dense_2 (128→64)", "128×64 + 64", "8,256"],
        ["Dense_3 (64→32)", "64×32 + 32", "2,080"],
        ["Output (32→8)", "32×8 + 8", "264"],
        ["+ BatchNorm", "—", "≈384"],
        ["TOPLAM", "", "≈ 18,152"],
    ]
)

doc.add_heading("9.4 Toplam Sistem", level=2)

make_simple_table(doc,
    ["Model", "Parametre", "Eğitim Verisi"],
    [
        ["Autoencoder (DT)", "≈ 11,646", "Sadece benign trafik"],
        ["Random Forest", "100 ağaç", "Tüm eğitim verisi"],
        ["Conv1D-LSTM", "≈ 55,336", "Tüm eğitim verisi (sekanslar)"],
        ["Hibrit Füzyon MLP", "≈ 18,152", "AE + ConvLSTM çıkışları"],
        ["TOPLAM", "≈ 85,134 + RF", "—"],
    ]
)


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  10. END-TO-END VERİ AKIŞI                                              ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

doc.add_page_break()
doc.add_heading("10. End-to-End Matematiksel Veri Akışı", level=1)

doc.add_paragraph(
    "Tek bir test örneği (xᵢ ∈ ℝ⁴⁶) için tam pipeline akışı:"
)

steps = [
    ("Adım 1 — Normalizasyon",
     "x'ᵢ = (xᵢ − μ) / σ",
     "Ham trafik verisi standartlaştırılır."),
    ("Adım 2 — Pencere Oluşturma",
     "X_seq = [x'_{i-9}, x'_{i-8}, ..., x'_i] ∈ ℝ^{10×46}",
     "Son 10 zaman adımı bir pencere oluşturur."),
    ("Adım 3 — AE Rekonstrüksiyon",
     "x̂ᵢ = f_dec(f_enc(x'ᵢ))",
     "Son zaman adımı AE'den geçirilir."),
    ("Adım 4 — AE Özellik Hataları",
     "e_feat = [(x'ᵢ₁ − x̂ᵢ₁)², ..., (x'ᵢ₄₆ − x̂ᵢ₄₆)²] ∈ ℝ⁴⁶",
     "46 boyutlu özellik hata vektörü."),
    ("Adım 5 — AE Global Skor",
     "e_global = mean(e_feat) ∈ ℝ¹",
     "Skaler anomali skoru."),
    ("Adım 6 — ConvLSTM Tahmin",
     "p = Conv1D_LSTM(X_seq) ∈ ℝ⁸",
     "8 sınıf olasılık vektörü."),
    ("Adım 7 — Füzyon",
     "v = [e_feat ‖ e_global ‖ p] ∈ ℝ⁵⁵",
     "55 boyutlu birleşik vektör."),
    ("Adım 8 — Karar",
     "ŷ = Softmax(MLP(v)) ∈ ℝ⁸",
     "Final sınıflandırma kararı."),
    ("Adım 9 — RCA (anomali ise)",
     "kök_neden = argsort(e_feat)[-5:]",
     "En sapan 5 özellik belirlenir."),
]

for step_title, formula, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(step_title)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    add_formula(doc, formula)

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    run2 = p2.add_run("→ " + desc)
    run2.font.italic = True
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# ── KAYDET ──────────────────────────────────────────────────────────────────
output_path = os.path.join(
    os.path.expanduser("~/Desktop"),
    "CICIoT2023_Matematiksel_Model_Detaylari.docx"
)
doc.save(output_path)
print(f"[✓] Word belgesi kaydedildi → {output_path}")
